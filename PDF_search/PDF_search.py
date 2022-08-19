from pdfminer.high_level import extract_text
import tkinter as tk
from tkinter import EXTENDED, ttk, filedialog
import re
import os
import sys
import docx

class Search(tk.Tk):
    def __init__(self):
        super().__init__()

        self.lst = []

        #GUI Base
        self.title("Keyword Search")
        self.foldersearch = ttk.Frame(self, padding = "10")
        self.foldersearch.grid(column = 0, row = 0)
        self.keywordsearch = ttk.Frame(self)
        self.keywordsearch.grid(column = 0, row = 1)
        self.submit = ttk.Frame(self)
        self.submit.grid(column = 0, row = 2)

        #Folder Search
        self.folderbox = ttk.Label(self.foldersearch, text = "", borderwidth = 3, relief= "solid")
        self.folderbox.config(width = 50, background = "white")
        self.folderbox.grid(column = 0, row = 0)
        folderbtn = ttk.Button(self.foldersearch, text = "Select Folder", command = self.openfolder)
        folderbtn.grid(column = 1, row = 0)

        #Keyword Search
        self.kword = tk.StringVar()
        self.keyword_entry = ttk.Entry(self.keywordsearch, textvariable = self.kword)
        self.keyword_entry.grid(column = 0, row = 0)

        #Submit
        self.submitbtn = ttk.Button(self.submit, text = "Search", command = self.search)
        self.submitbtn.grid(column = 0, row = 0)

    def openfolder(self):
        self.open_folder = filedialog.askdirectory()
        self.folderbox.config(text = self.open_folder)

    def search(self):
        self.lst = []
        folder = os.listdir(self.open_folder)
        for filename in folder:
            file = self.open_folder + "\\" + filename
            if file.endswith(".pdf"):
                f = extract_text(file)
                self.keyword(filename, f)
            elif file.endswith(".doc") or file.endswith(".docx"):
                doc = docx.Document(file)
                texts = []
                for para in doc.paragraphs:
                    texts.append(para.text)
                fulltext = " ".join(texts)
                self.keyword(filename, fulltext)
        if len(self.lst) == 0:
            self.lst.append("No files found!")
        self.newwindow()

    def keyword(self, filename, txt):
        search = re.search(self.kword.get(), txt, re.IGNORECASE)
        if search:
            self.lst.append(filename)


    def newwindow(self):
        wind = tk.Toplevel(self)
        self.title("Files Containing Keyword")
        self.boxframe = ttk.Frame(wind, padding = "10", width = 200)
        self.boxframe.grid(column = 0, row = 0)

        filelist = tk.StringVar(value = self.lst)
        self.box = tk.Listbox(self.boxframe, listvariable = filelist, selectmode = EXTENDED)
        self.box.grid(column = 0, row = 0)
        openbtn = ttk.Button(self.boxframe, text = "Open Files", command = self.openfiles)
        openbtn.grid(column = 0, row = 1)

    def openfiles(self):
        files = self.box.curselection()
        for i in files:
            f = self.box.get(i)
            os.startfile(f)



def main(): 

    f_input = input("Folder to search: ").replace("/", "\\")
    folder = os.listdir(f_input)

    global text
    text = input("What keyword do you want to find? ")
    global lst
    lst = []

    if len(lst) == 0:
        print("Keyword not found anywhere!")
    else:
        print(lst)

    for filename in folder:
        if f_input.endswith("\\"):
            file = f_input + filename
        else:
            file = f_input + "\\" + filename
        if file.endswith(".pdf"):
            f = extract_text(file)
            keyword(filename, f)
        elif file.endswith(".doc") or file.endswith(".docx"):
            doc = docx.Document(file)
            texts = []
            for para in doc.paragraphs:
                texts.append(para.text)
            fulltext = " ".join(texts)
            keyword(filename, fulltext)

    choice = input("Open files? ").lower().strip()
    if choice == "yes" or choice == "y":
        for i in lst:
            os.startfile(i)
    else:
        sys.exit()

def keyword(filename, txt):
    search = re.search(text, txt, re.IGNORECASE)
    if search:
        lst.append(filename)
    else:
        return "no"

if __name__ == "__main__":
    app = Search()
    app.mainloop()






